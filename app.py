from flask import Flask, render_template, Response ,jsonify
import cv2
import numpy as np
import time
from werkzeug.utils import redirect
import docx
import os
import shutil
app = Flask(__name__)

log = {
    "cell phone" : 0,
    "person" : 0
}
log1 = {
    "cell phone" : 0,
    "person" : 0
}

def gen_frames():  
    count = 0
    os.mkdir("./images_log")
    mydoc = docx.Document()
    cap = cv2.VideoCapture(0)
    net = cv2.dnn.readNet("weights/yolov3-tiny.weights", "cfg/yolov3-tiny.cfg")
    classes = []
    with open("coco.names", "r") as f:
        classes = [line.strip() for line in f.readlines()]
    layer_names = net.getLayerNames()
    # outputlayers = [layer_names[i-1] for i in net.getUnconnectedOutLayers()]
    output_layers = [layer_names[i-1] for i in net.getUnconnectedOutLayers()]
    colors = np.random.uniform(0, 255, size=(len(classes), 3))

    font = cv2.FONT_HERSHEY_PLAIN
    starting_time = time.time()
    frame_id = 0
    currentframe = 0
    while True:
        success, frame = cap.read()  # read the camera frame
        if not success:
            break
        else:
            frame_id += 1

            height, width, channels = frame.shape
        
            # Detecting objects
            blob = cv2.dnn.blobFromImage(frame, 0.00392, (416, 416), (0, 0, 0), True, crop=False)
        
            net.setInput(blob)
            outs = net.forward(output_layers)
        
            # Showing informations on the screen
            class_ids = []
            confidences = []
            boxes = []
            for out in outs:
                for detection in out:
                    scores = detection[5:]
                    class_id = np.argmax(scores)
                    confidence = scores[class_id]
                    if confidence > 0.2:
                        # Object detected
                        center_x = int(detection[0] * width)
                        center_y = int(detection[1] * height)
                        w = int(detection[3] * width)
                        h = int(detection[3] * height)
        
                        # Rectangle coordinates
                        x = int(center_x - w / 1.8)
                        y = int(center_y - h / 1.8)
        
                        boxes.append([x, y, w, h])
                        confidences.append(float(confidence))
                        class_ids.append(class_id)
        
            indexes = cv2.dnn.NMSBoxes(boxes, confidences, 0.4, 0.3)
            if(len(boxes) == 0):
                log1["person"]+=1
                log["person"] = 0
            for i in range(len(boxes)):
                if i in indexes:
                    x, y, w, h = boxes[i]
                    label = str(classes[class_ids[i]])
                    if(label == "cell phone"):
                        log1[label]+=1
                        log[label] = 1
                    elif(label == "person"):
                        log[label] = 1
                    else:
                        log["cell phone"] = 0
                        log["person"] = 0
                    print(classes[class_ids[i]])
                    confidence = confidences[i]
                    color = colors[class_ids[i]]
                    cv2.rectangle(frame, (x, y), (x + w, y + h), color, 2)
                    cv2.putText(frame, label + " " + str(round(confidence, 2)), (x, y + 30), font, 2, color, 2)
        
        
        
            elapsed_time = time.time() - starting_time
            fps = frame_id / (elapsed_time)
            cv2.putText(frame, "FPS: " + str(round(fps, 2)), (10, 50), font, 2, (0, 0, 0), 3)
            ret, buffer = cv2.imencode('.jpg', frame)
            if ret:
                # if video is still left continue creating images
                
                if(log["cell phone"] == 1):
                    name = './images_log/' + str(currentframe) + '.jpg'
                    print ('Creating...' + name)
                    # writing the extracted images
                    cv2.imwrite(name, frame)
                    # increasing counter so that it will
                    # show how many frames are created
                    mydoc.add_heading("[ " + "Cell phone" + " detected !!!!!! ]", 2)
                    mydoc.add_picture(name, width=docx.shared.Inches(3), height=docx.shared.Inches(3))
                    mydoc.save("./Student1.docx")
                    log["cell phone"] = 0
                if(log["person"] == 0):
                    print(count)
                    if(count == 25):
                        name = './images_log/' + str(currentframe) + '.jpg'
                        print ('Creating...' + name)
                        # writing the extracted images
                        cv2.imwrite(name, frame)
                        # increasing counter so that it will
                        # show how many frames are created
                        mydoc.add_heading("[ " + "Person" + " not detected !!!!!! ]",2)
                        mydoc.add_picture(name, width=docx.shared.Inches(3), height=docx.shared.Inches(3))
                        mydoc.save("./Student1.docx")
                        count = 0
                    else:
                        count+=1
                else:
                    log["person"] = 1
                
                currentframe += 1
            frame = buffer.tobytes()
            yield (b'--frame\r\n'
                   b'Content-Type: image/jpeg\r\n\r\n' + frame + b'\r\n')  # concat frame one by one and show result               

@app.route('/',methods = ['GET','POST'])
def hello_world():
    path = "./images_log"
    # if os.path.exists(path):
    #     os.rmdir(path)
    shutil.rmtree(path, ignore_errors=True)
    return render_template("index.html");

@app.route('/camera',methods = ['GET','POST'])
def camera():
    return render_template("Camera.html");

@app.route('/student',methods = ['GET','POST'])
def student():
    return render_template("student.html");

@app.route('/teacher',methods = ['GET','POST'])
def teacher():
    return render_template("teacher.html");

@app.route('/_get_data/', methods=['POST'])
def _get_data():
    data_log = []
    for key,value in log1.items():
        data_log.append([key,value])
    # return jsonify({'data': render_template('response.html', data_log = data_log)})
    return jsonify({'data': data_log})


@app.route('/video_feed')
def video_feed():
    return Response(gen_frames(), mimetype='multipart/x-mixed-replace; boundary=frame')

if __name__ == "__main__":
    app.run(debug=True)